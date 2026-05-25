"""Build complete workSphere report as Word document."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT = Path(r"c:\Users\Khushi Jagga\Desktop\Training Project\workSphere-Full-Report.docx")
ASSETS = Path(r"c:\Users\Khushi Jagga\.cursor\projects\c-Users-Khushi-Jagga-Desktop-Training-Project\assets")


def add_title(doc, text, level=0):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
    elif level == 1:
        p = doc.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)


def add_body(doc, text):
    for para in text.strip().split("\n\n"):
        doc.add_paragraph(para.strip())


def add_table(doc, caption, headers, rows):
    doc.add_paragraph(caption).alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def add_figure(doc, caption, image_name=None):
    doc.add_paragraph(caption).alignment = WD_ALIGN_PARAGRAPH.CENTER
    img = ASSETS / image_name if image_name else None
    if img and img.exists():
        doc.add_picture(str(img), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Insert screenshot or diagram here]")
    doc.add_paragraph()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # COVER
    add_title(doc, "INDUSTRIAL TRAINING REPORT")
    add_title(doc, "on")
    add_title(doc, "WorkSphere: Team Collaboration Platform")
    add_title(doc, "at")
    add_title(doc, "NED Gaming Ltd")
    doc.add_paragraph()
    add_body(
        doc,
        """Submitted by
Khushi Jagga
2237210

Towards the fulfilment of the degree of
BACHELORS IN TECHNOLOGY
in
DEPARTMENT OF ELECTRONICS AND COMMUNICATION ENGINEERING
Chandigarh Engineering College, Mohali, Punjab

(2022 - 2026)

Industrial Training & Project Report
Name of Organisation: NED Gaming Ltd
Address: G 1/3 Model Town, Delhi 110009
Department of training: Full Stack Department
Software / Hardware: Software
Technology: Express.js, React, SQLite, JavaScript, Vite
Period of Training: January 2026 To April 2026
Date of submission: 22 May 2026""",
    )
    doc.add_page_break()

    # CERTIFICATE
    add_title(doc, "CERTIFICATE", 1)
    add_body(
        doc,
        """This is to certify that the report on WorkSphere (Team Collaboration Platform) has been prepared and submitted as an authentic software project document. It presents the complete design and implementation of a full-stack web solution that brings team messaging, shared work boards, planning views, video meetings, and smart assistance together.

The system is built with Express.js, SQLite, bcrypt authentication, React, JavaScript, and Vite, following a modular monolithic architecture with clearly separated API routes, collaboration logic, database models, validation, middleware, and configuration modules.

Submitted by: Khushi Jagga
Project Title: WorkSphere (Team Collaboration Platform)
Institution / Department: Chandigarh Engineering College, ECE Department
Training / Project Duration: January 2026 to April 2026

Signature of Guide / Supervisor: _________________________
Date: _________________________""",
    )
    doc.add_page_break()

    add_title(doc, "ACKNOWLEDGEMENT", 1)
    add_body(
        doc,
        """I would like to express my heartfelt thanks to my faculty members, project guide, and everyone who supported me during the development of WorkSphere (Team Collaboration Platform). Their suggestions and encouragement helped me shape this work into a practical, well-structured software solution.

This project gave me valuable hands-on experience in backend architecture, database design, REST API development, frontend user experience, authentication, real-time communication, deployment concepts, and team collaboration workflows.

Finally, I am grateful to my institution for providing the academic environment, resources, and motivation needed to complete this project in a disciplined and professional manner.

Khushi Jagga""",
    )
    doc.add_page_break()

    add_title(doc, "LIST OF FIGURES", 1)
    for line in [
        "Fig. 4.1 System Architecture — 19",
        "Fig. 4.2 Authentication Flow — 20",
        "Fig. 4.3 Collaboration Data Flow — 21",
        "Fig. 4.4 Meet Signaling and WebRTC Flow — 22",
        "Fig. 5.1 Database Relationship Diagram — 25",
        "Fig. 7.1 Frontend Structure and Design System — 35",
        "Fig. 7.2 Login and Signup Interface — 36",
        "Fig. 7.3 workSphere Chat — 37",
        "Fig. 7.4 Team Invite Flow — 38",
        "Fig. 7.5 Shared Boards and Backlog — 39",
        "Fig. 7.6 Meet Interface — 40",
        "Fig. 7.7 Whiteboard Interface — 41",
        "Fig. 7.8 AI Assistant and Product Tour — 42",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    add_title(doc, "LIST OF TABLES", 1)
    for line in [
        "Table 1.1 Company Profile Summary — 8",
        "Table 1.2 workSphere Application Modules — 9",
        "Table 2.1 Project Objectives — 13",
        "Table 2.2 Existing System and Proposed System — 14",
        "Table 3.1 Software Requirements — 15",
        "Table 3.2 Hardware Requirements — 15",
        "Table 3.3 Technologies Used — 16",
        "Table 5.1 Database Tables — 24",
        "Table 6.1 Backend Module Responsibilities — 28",
        "Table 8.1 REST API Groups — 43",
        "Table 10.1 Testing Summary — 47",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    add_title(doc, "TABLE OF CONTENTS", 1)
    for line in [
        "Certificate — 2",
        "Acknowledgement — 3",
        "List of Figures — 4",
        "List of Tables — 5",
        "Chapter One: Company Profile — 8-10",
        "Chapter Two: Introduction — 11-14",
        "Chapter Three: Requirement Analysis — 15-17",
        "Chapter Four: System Design — 18-22",
        "Chapter Five: Database Design — 23-26",
        "Chapter Six: Backend Implementation — 27-32",
        "Chapter Seven: Frontend Implementation — 33-42",
        "Chapter Eight: APIs and Integrations — 43",
        "Chapter Nine: Workflow and Methodology — 44-45",
        "Chapter Ten: Testing and Deployment — 46-47",
        "Chapter Eleven: Results, Limitations, and Future Scope — 48-51",
        "References — 52",
        "Appendix — 53",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    # CHAPTER 1
    add_title(doc, "CHAPTER ONE", 1)
    add_title(doc, "COMPANY PROFILE", 1)
    add_body(
        doc,
        """Company Overview
N.E.D Gaming Pvt. Ltd. is a technology-driven company operating in productivity software, esports performance tracking, SEO automation, and AI analysis solutions. The organization develops software products for esports players, gaming influencers, content creators, and tournament organizers.

One major product is the Senchou App for performance tracking and esports scoring. The company focuses on cloud-based and AI-powered software for the gaming industry.""",
    )
    add_table(
        doc,
        "Table 1.1 Company Profile Summary",
        ["Particular", "Details"],
        [
            ["Company Name", "NED Gaming Ltd"],
            ["Business Domain", "Productivity software, esports, AI solutions"],
            ["Target Users", "Esports players, creators, teams"],
            ["Website", "https://nedgaming.com/"],
            ["Location", "G 1/3 Model Town, Delhi 110009"],
        ],
    )
    add_table(
        doc,
        "Table 1.2 workSphere Application Modules",
        ["Module", "Purpose"],
        [
            ["workSphere Chat", "Team channels, messaging, invites, attachments"],
            ["Shared Boards", "Kanban-style work tracking"],
            ["Meet", "Video meetings with WebRTC"],
            ["Whiteboard", "Visual collaboration"],
            ["AI Assistant", "PM guidance and product help"],
            ["Authentication", "Secure signup and login"],
        ],
    )
    add_body(
        doc,
        """Relevance to the Project
The company profile is relevant because both the training environment and WorkSphere focus on software delivery, team coordination, and structured product development. WorkSphere applies similar engineering discipline to collaboration workflows.""",
    )
    doc.add_page_break()

    # CHAPTER 2
    add_title(doc, "CHAPTER TWO", 1)
    add_title(doc, "INTRODUCTION", 1)
    add_body(
        doc,
        """2.1 Introduction
Team collaboration tools usually handle either chat or task tracking, not both in one flow. WorkSphere (Team Collaboration Platform) combines team channels, shared boards, planning views, video meetings, whiteboard work, and an AI assistant in one workspace.

The system is meant for software and product teams who lose context when updates live in one app and delivery work in another. workSphere addresses this with team chat, Kanban-style boards, Meet, whiteboard, and guided assistance.

The build uses Express.js and SQLite on the backend, React, JavaScript, and Vite on the frontend, bcrypt for secure passwords, and WebSocket signaling for Meet.""",
    )
    add_figure(doc, "Technology Stack", "worksphere-tech-stack-logos.png")
    add_body(
        doc,
        """2.2 Project Overview
WorkSphere is a team collaboration and project-delivery web application. Users can create accounts, join team channels, send messages, invite teammates, upload attachments, join video meetings, use a whiteboard, and interact with an AI assistant.

The project joins communication with delivery tracking. The app is built with the backend and frontend hosted separately for easier scaling and maintenance.

2.4 Problem Statement
Many teams lose time through small coordination gaps. Traditional tools show messages or tasks separately but do not show whether the team is aligned or what to do next. The problem is to build one application that connects communication and delivery work.""",
    )
    add_table(
        doc,
        "Table 2.1 Project Objectives",
        ["Objective", "Description"],
        [
            ["Team communication", "Channels, messages, invites, attachments"],
            ["Delivery analysis", "Boards, backlog, planning views"],
            ["Collaboration guidance", "AI assistant for PM and product help"],
            ["Team awareness", "Unified workspace for all roles"],
            ["Production readiness", "Modular architecture, APIs, testing, deployment"],
        ],
    )
    add_table(
        doc,
        "Table 2.2 Existing System and Proposed System",
        ["Aspect", "Existing System", "Proposed System"],
        [
            ["Team communication", "Standalone chat apps", "workSphere chat with invites"],
            ["Delivery visibility", "Separate board tools", "Shared Kanban and planning"],
            ["Meetings", "Basic meeting links", "Meet + whiteboard integrated"],
            ["Guidance", "External documentation", "AI assistant + product tour"],
            ["Architecture", "Many paid products", "Modular Express + React stack"],
        ],
    )
    doc.add_page_break()

    # CHAPTER 3
    add_title(doc, "CHAPTER THREE", 1)
    add_title(doc, "REQUIREMENT ANALYSIS", 1)
    add_table(
        doc,
        "Table 3.1 Software Requirements",
        ["Component", "Requirement"],
        [
            ["Backend", "Express.js (Node.js)"],
            ["Database", "SQLite"],
            ["Authentication", "bcrypt password hashing"],
            ["Frontend", "React, JavaScript, Vite"],
            ["Real-time", "WebSocket for Meet"],
            ["Testing", "ESLint, build checks, manual tests"],
            ["Deployment", "Vercel (frontend) + Node host (API)"],
        ],
    )
    add_table(
        doc,
        "Table 3.2 Hardware Requirements",
        ["Component", "Requirement"],
        [
            ["Processor", "Intel i3 / AMD Ryzen 3 or higher"],
            ["RAM", "4 GB minimum (8 GB recommended)"],
            ["Storage", "10 GB free"],
            ["Network", "Internet for API and Meet"],
            ["OS", "Windows 10/11, Linux, or macOS"],
        ],
    )
    add_table(
        doc,
        "Table 3.3 Technologies Used",
        ["Technology", "Use in Project"],
        [
            ["Express.js", "REST APIs and middleware"],
            ["SQLite", "Persistent collaboration data"],
            ["bcrypt", "Password security"],
            ["React + Vite", "Frontend UI"],
            ["WebSocket", "Meet signaling"],
            ["Multer", "File uploads"],
            ["OpenAI API", "AI assistant (optional)"],
        ],
    )
    add_body(
        doc,
        """3.3 Development and Deployment Environment
The project lives in the jira-ui folder. npm run dev starts the API on port 8787 and Vite on 5173. Production uses VITE_API_URL so the Vercel frontend can reach the deployed API.""",
    )
    doc.add_page_break()

    # CHAPTER 4
    add_title(doc, "CHAPTER FOUR", 1)
    add_title(doc, "SYSTEM DESIGN", 1)
    add_body(
        doc,
        """4.1 System Design Approach
The system follows a modular monolithic architecture. The API layer accepts HTTP requests, service logic contains business rules, data access handles SQLite operations, and models define database tables.

4.2 System Architecture
The architecture separates the React client, Express API routes, collaboration services, SQLite database, and Meet/WebSocket signaling.""",
    )
    add_figure(doc, "Fig. 4.1 System Architecture", "fig-6-1-express-backend-architecture.png")
    add_body(doc, """4.3 Authentication Flow
Users register or log in on the frontend. The backend checks passwords with bcrypt and returns user details. The browser stores the session for collaboration API calls.""")
    add_figure(doc, "Fig. 4.2 Authentication Flow", "fig-5-3-database-relationship-diagram.png")
    add_body(doc, """4.4 Data Processing Pipeline
Collaboration requests are validated, processed by the API, stored in SQLite, and returned to the frontend for chat display and assistant support.""")
    add_figure(doc, "Fig. 4.3 Collaboration Data Flow", "fig-6-2-api-service-repository-layers.png")
    doc.add_page_break()

    # CHAPTER 5
    add_title(doc, "CHAPTER FIVE", 1)
    add_title(doc, "DATABASE DESIGN", 1)
    add_body(
        doc,
        """5.1 Database Design Overview
The database is designed around registered users and team collaboration data. SQLite stores users, rooms, members, messages, and email invites in relational tables.""",
    )
    add_table(
        doc,
        "Table 5.1 Database Tables",
        ["Table", "Important Fields", "Purpose"],
        [
            ["users", "id, email, password_hash", "Accounts"],
            ["collab_rooms", "id, name, invite_token", "Channels"],
            ["collab_members", "room_id, email", "Membership"],
            ["collab_messages", "room_id, body", "Chat messages"],
            ["collab_email_invites", "accept_token, status", "Email invites"],
        ],
    )
    add_figure(doc, "Fig. 5.1 Database Tables", "fig-5-1-worksphere-database-tables.png")
    add_figure(doc, "Fig. 5.1 Database Relationship Diagram", "fig-5-3-database-relationship-diagram.png")
    add_body(
        doc,
        """5.4 Data Integrity and Migration Strategy
Schema updates run at server startup using CREATE TABLE and ensure functions. Validation, foreign keys, and UNIQUE constraints protect data integrity.""",
    )
    doc.add_page_break()

    # CHAPTER 6
    add_title(doc, "CHAPTER SIX", 1)
    add_title(doc, "BACKEND IMPLEMENTATION", 1)
    add_body(doc, """6.1 Backend Structure
The backend is implemented in Express.js with modular routes for authentication, collaboration, Meet signaling, and the AI assistant.""")
    add_figure(doc, "Fig. 6.1 Express Backend Architecture", "fig-6-1-express-backend-architecture.png")
    add_table(
        doc,
        "Table 6.1 Backend Module Responsibilities",
        ["Layer", "Responsibility"],
        [
            ["API endpoints", "HTTP requests and responses"],
            ["Validation", "Input and membership checks"],
            ["Services", "Rooms, messages, invites"],
            ["Data access", "SQLite queries"],
            ["Middleware", "CORS, JSON, uploads"],
        ],
    )
    add_figure(doc, "Fig. 6.2 API Layers", "fig-6-2-api-service-repository-layers.png")
    add_body(
        doc,
        """6.2 Collaboration Module — Channel creation, messaging, invites, attachments.
6.3 Meet Module — WebSocket signaling for WebRTC.
6.4 AI Assistant — OpenAI-powered chat endpoint.""",
    )
    doc.add_page_break()

    # CHAPTER 7
    add_title(doc, "CHAPTER SEVEN", 1)
    add_title(doc, "FRONTEND IMPLEMENTATION", 1)
    add_body(
        doc,
        """7.1 Frontend Structure and Design System
The frontend uses React, JavaScript, and Vite. Pages are organized by feature with shared components for navigation, chat, and theme support.""",
    )
    add_figure(doc, "Fig. 7.1 Frontend Structure and Design System", "fig-7-1-frontend-design-system-worksphere.png")
    for cap in [
        ("Fig. 7.2 Login and Signup Interface", None),
        ("Fig. 7.3 workSphere Chat", None),
        ("Fig. 7.4 Team Invite Flow", None),
        ("Fig. 7.5 Shared Boards", None),
        ("Fig. 7.6 Meet Interface", None),
        ("Fig. 7.7 Whiteboard", None),
        ("Fig. 7.8 AI Assistant", None),
    ]:
        add_figure(doc, cap[0], cap[1])
    doc.add_page_break()

    # CHAPTER 8-11
    add_title(doc, "CHAPTER EIGHT", 1)
    add_title(doc, "APIs AND INTEGRATIONS", 1)
    add_table(
        doc,
        "Table 8.1 REST API Groups",
        ["API Group", "Purpose"],
        [
            ["Auth", "/api/signup, /api/login"],
            ["Collaboration", "/api/collab/rooms, messages, invites"],
            ["Uploads", "File attachments"],
            ["Meet", "/meet-ws WebSocket"],
            ["Assistant", "/api/assistant/chat"],
        ],
    )
    doc.add_page_break()

    add_title(doc, "CHAPTER NINE", 1)
    add_title(doc, "WORKFLOW AND METHODOLOGY", 1)
    add_body(doc, """Incremental development: Express + SQLite + React first, then chat, Meet, whiteboard, and AI assistant. Local testing with npm run dev.""")
    doc.add_page_break()

    add_title(doc, "CHAPTER TEN", 1)
    add_title(doc, "TESTING AND DEPLOYMENT", 1)
    add_table(
        doc,
        "Table 10.1 Testing Summary",
        ["Check", "Result"],
        [
            ["ESLint", "Passed"],
            ["Production build", "Passed"],
            ["Signup / Login", "Passed"],
            ["Create channel", "Passed"],
            ["Send message", "Passed"],
            ["Meet UI", "Passed"],
        ],
    )
    add_body(doc, """Deployment: Frontend on Vercel, backend on Node.js host. Set VITE_API_URL for production chat.""")
    doc.add_page_break()

    add_title(doc, "CHAPTER ELEVEN", 1)
    add_title(doc, "RESULTS, LIMITATIONS, AND FUTURE SCOPE", 1)
    add_body(
        doc,
        """Results: WorkSphere delivers integrated team collaboration with chat, boards, Meet, whiteboard, and AI help.

Limitations: Not full enterprise replacement; backend required for chat on Vercel.

Future Scope: PostgreSQL, JWT, real-time sync, mobile app, GitHub integration.

Conclusion: The project demonstrates a complete full-stack collaboration platform using React, Express, SQLite, bcrypt, and WebSocket technologies.""",
    )
    doc.add_page_break()

    add_title(doc, "REFERENCES AND BIBLIOGRAPHY", 1)
    for ref in [
        "Express.js — https://expressjs.com/",
        "SQLite — https://www.sqlite.org/",
        "React — https://react.dev/",
        "Vite — https://vite.dev/",
        "Node.js — https://nodejs.org/",
        "WebRTC — https://webrtc.org/",
    ]:
        doc.add_paragraph(ref)

    add_title(doc, "APPENDIX", 1)
    add_body(
        doc,
        """POST /api/collab/rooms — create channel
POST /api/login — authenticate user
GET /api/collab/rooms/mine — list user channels
Pattern: API Route → Validation → Service → SQLite → JSON Response""",
    )

    doc.save(str(OUT))
    print(f"Created: {OUT}")


if __name__ == "__main__":
    build()
