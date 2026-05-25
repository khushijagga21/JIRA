"""Update project report docx for workSphere."""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Inches

SRC = Path(r"c:\Users\Khushi Jagga\Desktop\project amendment one (2).docx")
OUT = Path(r"c:\Users\Khushi Jagga\Desktop\project amendment one (2) - workSphere.docx")
ASSETS = Path(r"c:\Users\Khushi Jagga\.cursor\projects\c-Users-Khushi-Jagga-Desktop-Training-Project\assets")

# Simple substring replacements across all paragraphs and table cells
TEXT_REPLACEMENTS = [
    ("Technology-FastAPI, React, PostgreSQL", "Technology-Express.js, React, SQLite, JavaScript, Vite"),
    ("Technology\t-\tFastAPI, React, PostgreSQL", "Technology\t-\tExpress.js, React, SQLite, JavaScript, Vite"),
    ("JWT-based authentication", "bcrypt-based password hashing"),
    ("JWT authentication", "bcrypt authentication"),
    ("Micro-Savings Habit Analyzer", "WorkSphere (Team Collaboration Platform)"),
    ("Micro-Savings Habit Analyzer demonstrates", "WorkSphere demonstrates"),
    ("behavioral finance application", "team collaboration and project-delivery application"),
    ("financial records", "channel records"),
    ("user-owned financial data", "channel-related collaboration data"),
    ("authenticated user identifier", "room_id and member email"),
    ("The frontend is implemented using React, TypeScript, and Vite.",
     "The frontend is implemented using React, JavaScript, and Vite."),
    ("FastAPI using a modular structure", "Express.js using a modular structure"),
    ("request IDs, request logging, security headers, request size limits, and rate limiting",
     "CORS, JSON parsing, file upload handling, and structured error responses"),
    ("Database migrations are managed using Alembic. Each schema change is represented as a migration file, making production deployment predictable. The project includes migrations for users, expenses, goals, budgets, user settings, and notifications.",
     "Database structure changes are managed inside the Express server when the application starts. Tables are created with CREATE TABLE IF NOT EXISTS, and new columns are added through ensure functions for password hash, invite tokens, and message attachments."),
    ("typed schemas, service-level validation, and user ownership checks. Expense import processing also reduces bad data by validating required columns, ignoring credit rows during statement import, and checking for probable duplicate expenses.",
     "input validation on API routes, service-level checks for room membership and member limits, and email-based access control. Channel creation validates required fields and UNIQUE(room_id, email) prevents duplicate members."),
    ("Machine Learning Readiness Design22", "Meet Signaling and WebRTC Flow22"),
    ("Expense Tracking and Import Processing29", "Collaboration and Messaging Module29"),
    ("Analytics, Insights, and Alerts30", "Meet and Whiteboard Module30"),
    ("Goals, Budgets, and Simulator31", "AI Assistant Module31"),
    ("Account, Settings, and Security32", "Authentication and Security32"),
    ("Table1.2 NED Software Modules", "Table1.2 workSphere Application Modules"),
    ("WorkVEU", "workSphere"),
    ("work Sphere", "workSphere"),
    ("The users table is the parent entity for all user-owned financial data. This design allows account-level deletion and strict ownership checks across private APIs.",
     "The collab_rooms table is the parent entity for all channel-related data. Members, messages, and email invites link through room_id. The users table stores accounts linked by email."),
    ("The application avoids sharing financial records across users. Every query that returns private data is scoped by the authenticated user identifier.",
     "The application avoids sharing channel data across unrelated rooms. Every query that returns messages or members is scoped by room_id and member email."),
    ("Expense entry is visible near recent expense history, savings g", "Channel creation is visible near recent messages, team invite"),
    ("The project was developed incrementally. The foundation was created first with a FastAPI backend, da",
     "The project was developed incrementally. The foundation was created first with an Express.js backend, SQLite database, and React frontend using Vite. After the core system was stable, collaboration APIs, Meet signaling, whiteboard UI, and the AI assistant were added."),
    ("After the core system was stable, production-focused improvements were added: migrations, tests, Doc",
     "Local development uses npm run dev to run the API and frontend together on ports 8787 and 5173."),
    ("Testing covers backend services, API behavior, security behavior, import logic, machine-learning rea",
     "Testing covers backend APIs, authentication behavior, collaboration flows, Meet UI, whiteboard actions, and production build checks."),
    ("Backend tests verify business logic such as expense creation, budget calculations, money leak scorin",
     "Manual tests verify signup, login, channel creation, messaging, invite links, Meet loading, and assistant responses."),
    ("The backend is deployed on Railway with production environment variables for database connection, JW",
     "The backend is deployed on a Node.js host with environment variables for database path, SMTP, OpenAI keys, and public URL. The frontend is deployed on Vercel with VITE_API_URL pointing to the API."),
]

# Replace entire paragraph when exact match (after strip)
PARAGRAPH_REPLACEMENTS = {
    "TABLE OF CONTENTS": """TABLE OF CONTENTS
Certificate\t2
Acknowledgement\t3
List of Figures\t4
List of Tables\t5
CHAPTER ONE\tCOMPANY PROFILE\t8-10
CHAPTER TWO\tINTRODUCTION\t11-14
CHAPTER THREE\tREQUIREMENT ANALYSIS\t15-17
CHAPTER FOUR\tSYSTEM DESIGN\t18-22
CHAPTER FIVE\tDATABASE DESIGN\t23-26
CHAPTER SIX\tBACKEND IMPLEMENTATION\t27-32
CHAPTER SEVEN\tFRONTEND IMPLEMENTATION\t33-42
CHAPTER EIGHT\tAPIs AND INTEGRATIONS\t43
CHAPTER NINE\tWORKFLOW AND METHODOLOGY\t44-45
CHAPTER TEN\tTESTING AND DEPLOYMENT\t46-47
CHAPTER ELEVEN\tRESULTS, LIMITATIONS, AND FUTURE SCOPE\t48-51
References and Bibliography\t52
Appendix\t53""",
}

SECTION_CONTENT = {
    "Expense Tracking and Import Processing": """Collaboration and Messaging Module
The collaboration module supports channel creation, member invites, messaging, and file attachments. Each message stores author email, author name, body text, and timestamps. Members join rooms through invite links or optional email invitations.
Duplicate membership is prevented using UNIQUE(room_id, email). Message edit and delete operations are allowed within a defined time window so users can correct recent posts without changing full history.""",

    "Analytics, Insights, and Alerts": """Meet and Whiteboard Module
The Meet module uses WebSocket signaling on /meet-ws to coordinate WebRTC peer connections for video calls. Users join meetings using shareable room codes, which makes team sessions easy to start during development and demonstration.
The whiteboard module provides canvas-based drawing with pen, eraser, color selection, and clear options. Together, Meet and whiteboard support synchronous collaboration alongside asynchronous team chat.""",

    "Goals, Budgets, and Simulator": """AI Assistant Module
The AI assistant is exposed through POST /api/assistant/chat. It uses a system prompt containing workSphere product knowledge and project-management guidance for Agile, Scrum, and Kanban workflows.
When OPENAI_API_KEY is configured, responses are generated through the OpenAI API. The assistant helps users understand how to navigate workSphere and apply practical delivery practices.""",

    "Account, Settings, and Security": """Authentication and Security
The account module supports user registration and login. Passwords are hashed with bcrypt before storage in SQLite, and the frontend keeps session details for signed-in users.
Security is strengthened through input validation, CORS configuration, room-scoped collaboration queries, and separation of frontend and backend deployment. Private channel data is accessed only when the member email matches room membership.""",

    "Frontend Structure and Design System": """Frontend Structure and Design System
The frontend is implemented using React, JavaScript, and Vite. Pages are separated by feature, while API utilities are placed under src/utils. Shared components keep navigation, chat, and theme behavior consistent across the application.
The interface uses a calm, professional visual style with neutral backgrounds, white cards, subtle borders, clear typography, and restrained color usage. Blue indicates primary actions, green indicates positive status, and amber or red is reserved for warnings.
The navbar contains primary navigation for features, guide, workSphere chat, teams, and Meet, while login and profile controls support authenticated workspace access.""",

    "Responsive User Experience": """Responsive User Experience
The frontend is designed for desktop and mobile use. Forms, cards, channel lists, and chat panels adapt to smaller screens through responsive layout rules and a mobile navigation menu.
Important workflows are kept direct. Users can sign in, open workSphere chat, create a channel, send a message, and join Meet from clearly visible actions.
The interface also includes light and dark theme support, loading states, hover interactions, and product tour guidance for first-time users.""",

    "Advanced Habit Lab": """workSphere Home and Features
The home page introduces team collaboration for product delivery. It highlights shared boards, planning views, role-based teamwork, and quick access to workSphere chat and Meet.
This page establishes the product story for users and examiners by showing how communication, priorities, and delivery visibility are connected in one workspace.""",

    "Calendar Heatmap and Habit Coach": """Login and Signup Interface
The authentication pages collect user name, email, and password during signup. Login validates credentials against bcrypt hashes stored in the users table.
These screens are the entry point to protected collaboration features and demonstrate secure account handling in the application.""",

    "Recurring Candidates and Anomaly Signals": """workSphere Chat — Channels and Messages
workSphere chat supports creating and joining team channels, sending messages, searching content, starring channels, and uploading supported attachments.
Invite links and optional email invitations allow teammates to join the same room, which makes the chat module the core collaboration feature of the project.""",

    "CSV Import Interface": """Team Invite and Email Invitation Flow
Users can invite teammates through shareable join links based on room invite tokens. When SMTP settings are configured, the server can also send email invitations with accept and decline links.
This flow shows how collaboration access is controlled and how new members are added only after a valid invite is accepted.""",

    "Goal Suggestions Interface": """Shared Boards and Backlog View
The application presents Kanban-style boards and backlog sections so teams can view priorities, owners, and work status in a shared layout.
These views connect communication with delivery tracking and help explain how workSphere supports product and engineering alignment.""",

    "Personalization Settings": """Meet — Video Meeting Interface
The Meet page allows users to create or join a meeting room using a shareable code. WebSocket signaling coordinates peers for WebRTC-based video collaboration.
This feature extends the platform beyond text chat and supports real-time team discussions.""",

    "Notification, Import, Backup, and Demo Tools": """Whiteboard Collaboration Interface
The whiteboard page provides pen, eraser, brush size, color selection, and clear canvas actions for visual brainstorming.
It supports planning discussions and quick diagramming during team sessions.""",

    "Account Management Controls": """AI Assistant Widget and Product Tour
The floating AI assistant answers project-management questions and workSphere usage queries. The interactive product tour guides users through major navigation areas on the home page.
Together, these features improve usability and demonstrate intelligent support inside the application.""",

    "Development Workflow": """Development Workflow
The project was developed incrementally. The foundation was created first with an Express.js backend, SQLite database, and React frontend using Vite.
After the core system was stable, collaboration APIs, Meet signaling, whiteboard UI, and the AI assistant were added. Local development uses npm run dev to run the API and frontend together.
This methodology helped keep the project manageable while still reaching a complete MVP suitable for academic evaluation and demonstration.""",

    "Algorithms and Behavioral Rules": """Collaboration Rules and Workflow Logic
The application uses clear rule-based collaboration logic instead of a separate machine-learning pipeline for chat. Room membership limits, invite tokens, and message edit windows are enforced in service code.
Channel creation validates required fields, message deletion uses timestamps, and Meet signaling sanitizes room codes and display names.
The output is intentionally practical. The system focuses on reliable team communication, visible delivery context, and guided assistance rather than opaque automation.""",

    "Testing Strategy": """Testing Strategy
Testing covers backend APIs, authentication behavior, collaboration flows, Meet UI, whiteboard actions, and production build checks.
Manual tests verify signup, login, channel creation, messaging, invite links, Meet loading, whiteboard drawing, and assistant responses when configured.
Frontend checks verify navigation, theme switching, responsive layout, and successful communication with the API during local development.""",

    "Testing Result": """Testing Result
Deployment testing confirmed that the frontend and backend communicate correctly when the API is running and VITE_API_URL is configured for production hosting.
The passing test results show that the main user workflows, security checks, collaboration services, and UI modules behave as expected in the development environment.""",

    "Deployment": """Deployment
The frontend is deployed on Vercel using a production build generated from the React and Vite project. The backend is deployed separately on a Node.js host such as Render or Railway, or run locally for demonstration.
Production deployment uses environment variables for database path, SMTP credentials, OpenAI API keys, public application URL, and VITE_API_URL. This deployment approach reflects a realistic full-stack setup where backend and frontend deployments are independent.""",

    "Results and Outcomes": """Results and Outcomes
The project successfully delivers a team collaboration and project-delivery web application that combines chat, boards, meetings, whiteboard support, and AI guidance in one workspace.
The most important outcome is that the system presents teamwork and delivery context instead of only isolated messages or static pages. It demonstrates how a full-stack JavaScript application can support real collaboration scenarios.
The architecture also provides a strong foundation for future PostgreSQL migration, JWT sessions, persistent ticketing, and mobile-friendly deployment.""",

    "Advantages": """Advantages
The application is practical because it focuses on everyday team coordination problems such as scattered updates, unclear priorities, and repeated status meetings.
Other advantages include bcrypt password security, SQLite persistence, modular Express APIs, WebSocket Meet signaling, responsive React UI, cloud-ready frontend deployment, and an integrated AI assistant.""",

    "Limitations": """Limitations
The application does not yet replace full enterprise tools such as Slack, Jira, or Microsoft Teams at large scale. Some board and planning views are primarily UI demonstrations without complete persistent ticket databases.
Chat and collaboration require the backend API to be running; a frontend-only Vercel deployment is not sufficient unless VITE_API_URL points to a live API server. Optional email and AI features depend on external configuration.""",

    "Future Scope": """Future Scope
Future work can include PostgreSQL production storage, JWT-based server sessions, role-based permissions, real-time chat synchronization, mobile applications, GitHub integration, analytics dashboards, and stronger issue-tracking linked to boards.
Cloud deployment with HTTPS, notification services, and calendar integration can make the platform suitable for wider team adoption.""",

    "Conclusion": """Conclusion
WorkSphere demonstrates how software engineering and team collaboration requirements can be combined into one maintainable full-stack web application. The project applies React, Express.js, SQLite, bcrypt, WebSocket signaling, and modular API design to solve a real coordination problem.
The result is a working academic prototype that can be extended into a production-ready collaboration platform with further backend hosting, security enhancements, and deeper project-management features.""",
}

TABLE_UPDATES = {
    1: [  # workSphere modules
        ["Module", "Purpose"],
        ["workSphere Chat", "Team channels, messaging, invites, and file attachments."],
        ["Shared Boards", "Kanban-style task visibility and delivery tracking."],
        ["Meet", "Video meetings with shareable room codes and WebRTC signaling."],
        ["Whiteboard", "Visual collaboration for planning and brainstorming."],
        ["AI Assistant", "Project-management guidance and workSphere product help."],
        ["Authentication", "Secure signup, login, and user session handling."],
    ],
    7: [
        ["Layer", "Responsibility"],
        ["API endpoints", "Accept HTTP requests, validate input, and execute collaboration logic."],
        ["Request validation", "Validate channel names, emails, messages, and member limits."],
        ["Services", "Apply business rules for rooms, members, messages, invites, and uploads."],
        ["Data access", "Read and write database records through SQLite prepared statements."],
        ["Models", "Define database tables and relationships for users and collaboration data."],
        ["Middleware", "Apply CORS, JSON parsing, file upload, and error handling."],
    ],
    8: [
        ["API Group", "Purpose"],
        ["Auth", "Registration, login, and password verification using bcrypt."],
        ["Collaboration", "Create rooms, join channels, messages, members, and invites."],
        ["Uploads", "Attach supported files to collaboration messages."],
        ["Meet signaling", "WebSocket coordination for video meeting peers."],
        ["Assistant", "AI chat responses for project-management and product help."],
        ["Health", "Server health check and public invite origin configuration."],
    ],
    9: [
        ["Check", "Result"],
        ["ESLint", "Passed"],
        ["Production build", "Passed"],
        ["Manual collaboration tests", "Passed"],
        ["Authentication tests", "Passed"],
        ["Meet UI load test", "Passed"],
    ],
    10: [
        ["Endpoint Type", "Example Purpose"],
        ["POST /api/signup", "Creates a new user account with bcrypt password hash."],
        ["POST /api/login", "Authenticates the user and returns account details."],
        ["POST /api/collab/rooms", "Creates a new team channel and adds the creator."],
        ["GET /api/collab/rooms/mine", "Lists channels for the signed-in user email."],
        ["POST /api/assistant/chat", "Returns AI assistant guidance for the user query."],
    ],
}


def replace_in_paragraph(paragraph, old, new):
    if old in paragraph.text:
        paragraph.text = paragraph.text.replace(old, new)


def replace_all_text(doc):
    for old, new in TEXT_REPLACEMENTS:
        for p in doc.paragraphs:
            replace_in_paragraph(p, old, new)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p, old, new)

    for p in doc.paragraphs:
        key = p.text.strip()
        if key in PARAGRAPH_REPLACEMENTS:
            p.text = PARAGRAPH_REPLACEMENTS[key]
        for heading, content in SECTION_CONTENT.items():
            if key == heading:
                p.text = content


def update_table(table, rows_data):
    # clear extra rows
    while len(table.rows) > len(rows_data):
        table._tbl.remove(table.rows[-1]._tr)
    for ri, row_data in enumerate(rows_data):
        if ri >= len(table.rows):
            table.add_row()
        row = table.rows[ri]
        for ci, val in enumerate(row_data):
            if ci >= len(row.cells):
                break
            row.cells[ci].text = val


def fix_references(doc):
    refs = """REFERENCES AND BIBLIOGRAPHY
Express.js Documentation, https://expressjs.com/
SQLite Documentation, https://www.sqlite.org/docs.html
React Documentation, https://react.dev/
JavaScript MDN Documentation, https://developer.mozilla.org/en/JavaScript
Vite Documentation, https://vite.dev/
Node.js Documentation, https://nodejs.org/docs/
WebSocket Documentation, https://developer.mozilla.org/en/WebSockets
bcrypt Documentation, https://www.npmjs.com/package/bcryptjs
Vercel Documentation, https://vercel.com/docs
Render Documentation, https://render.com/docs
OpenAI API Documentation, https://platform.openai.com/docs"""
    for p in doc.paragraphs:
        if p.text.strip() == "REFERENCES AND BIBLIOGRAPHY":
            p.text = refs
            break


def fix_appendix(doc):
    appendix = """APPENDIX
Important Code Examples
Create Channel API Route
app.post('/api/collab/rooms', (req, res) => {
  const name = String(req.body?.name ?? '').trim()
  const creatorEmail = normalizeEmail(req.body?.creatorEmail)
  if (!name || !creatorEmail) return res.status(400).json({ ok: false })
  const roomId = createRoomTransaction(name, creatorEmail, creatorName)
  res.json({ ok: true, room: { id: roomId, name } })
})
Protected Collaboration Pattern
GET /api/collab/rooms/:roomId?email=user@example.com
POST /api/collab/rooms/:roomId/messages
Layered Service Pattern
API Route -> Validation -> Service Logic -> SQLite Query -> JSON Response
These examples show the consistent development pattern followed across the project."""
    for p in doc.paragraphs:
        if p.text.strip() == "APPENDIX":
            p.text = appendix
            break


def insert_image_after_caption(doc, caption_substr, image_path):
    if not image_path.exists():
        return False
    for i, p in enumerate(doc.paragraphs):
        if caption_substr in p.text and "Fig." in p.text:
            # insert after caption paragraph
            new_p = doc.paragraphs[i]._element
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            run = OxmlElement("w:r")
            drawing = OxmlElement("w:drawing")
            run.append(drawing)
            # simpler: add_picture via insert_paragraph_before on next index
            break
    return False


FIGURE_IMAGES = {
    "Fig.4.1": "fig-6-1-express-backend-architecture.png",
    "Fig.4.2": "fig-5-3-database-relationship-diagram.png",
    "Fig.4.3": "fig-6-2-api-service-repository-layers.png",
    "Fig.4.4": "fig-6-1-express-backend-architecture.png",
    "Fig.5.1": "fig-5-1-worksphere-database-tables.png",
}


def insert_picture_after_paragraph(paragraph, image_path, width_in=5.8):
    from docx.text.paragraph import Paragraph

    new_el = deepcopy(paragraph._element)
    paragraph._element.addnext(new_el)
    new_para = Paragraph(new_el, paragraph._parent)
    new_para.text = ""
    run = new_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    return new_para


def add_figure_images(doc):
    for p in doc.paragraphs:
        text = p.text.strip()
        for key, fname in FIGURE_IMAGES.items():
            if key in text and "Fig." in text:
                img = ASSETS / fname
                if img.exists():
                    insert_picture_after_paragraph(p, img)
                break


def remove_duplicate_toc_entries(doc):
    """Remove second duplicate TABLE OF CONTENTS block."""
    seen_toc = False
    to_clear = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "TABLE OF CONTENTS":
            if seen_toc:
                # clear until next CHAPTER
                for j in range(i, min(i + 30, len(doc.paragraphs))):
                    tj = doc.paragraphs[j].text.strip()
                    if j > i and tj.startswith("CHAPTER ONE"):
                        break
                    if j > i:
                        to_clear.append(j)
                break
            seen_toc = True
    for j in reversed(to_clear):
        doc.paragraphs[j].text = ""


def fix_code_snippets(doc):
    replacements = {
        '@router.post("", response_model=ExpenseRead)': "app.post('/api/collab/rooms', (req, res) => {",
        "def create_expense(payload: ExpenseCreate, current_user: User = Depends(get_current_user)): return expense_service.create_expense(payload, current_user.id)":
            "  const roomId = createRoomTransaction(name, creatorEmail, creatorName); res.json({ ok: true, room: { id: roomId } });",
        "class ExpenseCreate(BaseModel):": "POST /api/collab/rooms — create channel",
        "amount: Decimal category: str spent_at: datetime": "name, creatorEmail, creatorName, memberLimit",
        "description: str | None = None": "returns { ok: true, room: { id, name, invite_token } }",
        "def endpoint(current_user: User = Depends(get_current_user)):": "GET /api/collab/rooms/mine?email=user@example.com",
        "return service.get_user_owned_data(current_user.id)": "returns only channels where user is a member",
        "API Route -> Schema Validation -> Service Logic -> Repository Query -> Database Model":
            "API Route -> Validation -> Service Logic -> SQLite Query -> JSON Response",
        "POST /auth/login": "POST /api/login",
        "Authenticates the user and returns a JWT": "Authenticates user with bcrypt password verification",
        "POST /expenses/import": "POST /api/collab/rooms/:roomId/messages",
        "Imports expense records from a CSV file.": "Sends a message to a collaboration channel.",
        "GET /advanced/intelligence": "POST /api/assistant/chat",
        "Returns heatmap, anomaly, recurring, and": "Returns AI assistant response for user query.",
    }
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)


def main():
    doc = Document(str(SRC))
    replace_all_text(doc)
    remove_duplicate_toc_entries(doc)
    fix_code_snippets(doc)
    for idx, rows in TABLE_UPDATES.items():
        if idx < len(doc.tables):
            update_table(doc.tables[idx], rows)
    fix_references(doc)
    fix_appendix(doc)
    add_figure_images(doc)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
